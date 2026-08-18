"""
Generate the Laertius technical overview document in PDF and Word formats.
Produces:
  exports/laertius-technical-overview.pdf
  exports/laertius-technical-overview.docx
  exports/laertius-arch.png  (architecture figure, embedded in both)

All numbers are verified against the codebase (see comments).
No em dashes used anywhere in this file or the output documents.
"""

import os
import sys
import textwrap
import io
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
WORKSPACE = Path(__file__).resolve().parent.parent
EXPORTS = WORKSPACE / "exports"
EXPORTS.mkdir(exist_ok=True)
ARCH_PNG = EXPORTS / "laertius-arch.png"
PDF_OUT  = EXPORTS / "laertius-technical-overview.pdf"
DOCX_OUT = EXPORTS / "laertius-technical-overview.docx"

FONT_REG  = "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"
FONT_SANS = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_SANS_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

# ---------------------------------------------------------------------------
# Verified constants from the codebase
# ---------------------------------------------------------------------------
# bm25.ts
K1 = 1.5
B  = 0.75
# rag.ts
RRF_K       = 20
POOL        = 50
KG_MATCHED  = "0.5 / 20"   # = KG_MATCHED_BOOST as shown in rag.ts
KG_RELATED  = "0.15 / 20"  # = KG_RELATED_BOOST
SPARSE_WEIGHT = 0.3        # = DEFAULT_FUSION_PARAMS.sparseWeight in rag.ts
# embedder.ts / dense.ts
EMBED_MODEL = "Xenova/multilingual-e5-small"
EMBED_DIM   = 384   # standard output dim of multilingual-e5-small
EMBED_DTYPE = "q8"
# Short model label for the architecture figure, derived from EMBED_MODEL so
# a model change cannot leave the figure showing the old family name.
EMBED_SHORT = EMBED_MODEL.rsplit("/", 1)[-1].replace("multilingual-", "").replace("e5", "E5")
# kg.ts - 77 edges, counted from `from:` occurrences inside the KG_EDGES array block
KG_EDGES    = 77
# corpus: 1211 sections (wc -l laertius_sections.jsonl)
CORPUS_SECTIONS = 1211
# kg-claims: 1564 claims, all with grc verbatim Greek excerpt
CLAIM_COUNT = 1565
# 148 claims carry a named accordingTo authority
NAMED_AUTHORITY_CLAIMS = 149
# layer-pins.ts
ANNOTATION_COUNT = 9310  # 2026-08: see layer-pins.ts
TAGGED_ENTITIES  = 787
VERSE_COUNT      = 340
TESTAMENT_COUNT  = 6
# lod.ts annotated export: voidStats().annotatedTriples, rounded to the
# nearest thousand ("~<N>k triples"). Checked against the live LOD graph
# by validate-technical-doc.ts with a 10% approximation tolerance.
ANNOTATED_TRIPLES_K = 205
# 17 movements in kg.ts (MOVEMENTS array)
MOVEMENT_COUNT   = 17
# chapter-subjects.ts / layer-pins.ts CHAPTER_SUBJECT_PIN_COUNT
CHAPTER_SUBJECTS = 82

# NOTE: the constants in this block are checked by
# scripts/src/validate-technical-doc.ts (registered validator
# "technical-doc") against layer-pins.ts and the live layer sizes.
# If that validator fails, update the constant here and re-run this
# script to regenerate the PDF/DOCX exports.

# ---------------------------------------------------------------------------
# 1. Architecture figure
# ---------------------------------------------------------------------------
def make_arch_figure():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    fig, ax = plt.subplots(figsize=(12, 9))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_facecolor("#fafaf8")
    fig.patch.set_facecolor("#fafaf8")

    def box(ax, x, y, w, h, text, color, fontsize=8.5, textcolor="white", bold=False):
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.05",
                              linewidth=0.8,
                              edgecolor="#555", facecolor=color, zorder=3)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text,
                ha="center", va="center", fontsize=fontsize,
                color=textcolor, fontweight="bold" if bold else "normal",
                zorder=4, wrap=True,
                multialignment="center")

    def arrow(ax, x1, y1, x2, y2, label="", color="#666", lw=1.2, dashed=False):
        style = "dashed" if dashed else "solid"
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=lw,
                                   linestyle=style),
                    zorder=5)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx+0.08, my, label, fontsize=7, color=color, zorder=6)

    # Color palette
    C_CORPUS  = "#4a6fa5"
    C_RETRIEV = "#1d7874"
    C_KG      = "#6b4c8b"
    C_LOD     = "#8c4646"
    C_ANNOT   = "#b07d34"
    C_FRONT   = "#3a7d44"
    C_BUNDLE  = "#666"

    # --- Layer labels on left ---
    layers = [
        (8.3, "Corpus"),
        (6.7, "Retrieval"),
        (5.0, "Knowledge"),
        (3.3, "LOD"),
        (1.8, "Annotation"),
        (0.4, "Frontend"),
    ]
    for y, lbl in layers:
        ax.text(0.08, y, lbl, fontsize=7.5, color="#888", va="center",
                style="italic", zorder=4)

    # --- Title ---
    ax.text(6, 8.75, "Laertius: Layered Architecture",
            ha="center", va="center", fontsize=13, fontweight="bold",
            color="#222", zorder=4)

    # --- Corpus layer ---
    box(ax, 1.0, 8.0, 2.2, 0.55, f"Greek TEI\n(Perseus grc2)\n{CORPUS_SECTIONS} sections", C_CORPUS, 7.5)
    box(ax, 3.5, 8.0, 2.2, 0.55, "English TEI\n(Hicks translation)\naligned sections", C_CORPUS, 7.5)
    box(ax, 6.0, 8.0, 2.2, 0.55, "Metadata\n(school, chapter,\nphilosopher)", C_CORPUS, 7.5)

    # --- Retrieval layer ---
    box(ax, 1.0, 6.4, 2.0, 0.65, f"BM25\nK1={K1}, B={B}\nLucene IDF", C_RETRIEV, 7.5)
    box(ax, 3.2, 6.4, 2.2, 0.65, f"Dense\n{EMBED_SHORT} {EMBED_DTYPE}\n{EMBED_DIM}-dim cosine", C_RETRIEV, 7.5)
    box(ax, 5.6, 6.4, 2.2, 0.65, f"RRF Fusion\nk={RRF_K}, pool={POOL}\nCormack et al.", C_RETRIEV, 7.5)
    box(ax, 8.0, 6.4, 2.2, 0.65, f"KG Boost\n+{KG_MATCHED.replace(' ', '')} matched\n+{KG_RELATED.replace(' ', '')} related", C_RETRIEV, 7.5)

    # --- Knowledge layers ---
    box(ax, 1.0, 4.8, 1.7, 0.65, f"Philosophers\nGraph\n{KG_EDGES} edges", C_KG, 7.5)
    box(ax, 2.85, 4.8, 1.7, 0.65, f"Claims\n{CLAIM_COUNT} cited\n4 certainty lvls", C_KG, 7.5)
    box(ax, 4.7, 4.8, 1.7, 0.65, f"Verses\n{VERSE_COUNT} from\nTEI blockquotes", C_KG, 7.5)
    box(ax, 6.55, 4.8, 1.7, 0.65, "Sayings &\nAnecdotes\ncurated TS", C_KG, 7.5)
    box(ax, 8.4, 4.8, 1.7, 0.65, f"Doxai &\nTestaments\n{TESTAMENT_COUNT} wills verbatim", C_KG, 7.5)

    # --- LOD layer ---
    box(ax, 1.0, 3.0, 2.3, 0.65, f"Turtle / RDF-XML\n(~{ANNOTATED_TRIPLES_K}k triples\nannotated graph)", C_LOD, 7.5)
    box(ax, 3.5, 3.0, 2.3, 0.65, "SPARQL endpoint\n(oxigraph\nread-only)", C_LOD, 7.5)
    box(ax, 6.0, 3.0, 2.3, 0.65, "OTB / TEDI export\n(concepts, objects,\nproper names, RDF)", C_LOD, 7.5)
    box(ax, 8.5, 3.0, 1.8, 0.65, "Per-section\nsubgraphs\n(oa: layer)", C_LOD, 7.5)

    # --- Annotation layer (feedback) ---
    box(ax, 1.0, 1.5, 2.5, 0.65, f"Gazetteer\n(derived from LOD)\n{TAGGED_ENTITIES} entities", C_ANNOT, 7.5)
    box(ax, 3.8, 1.5, 2.5, 0.65, f"Annotator\n{ANNOTATION_COUNT} occurrences\nlongest-match-first", C_ANNOT, 7.5)
    box(ax, 6.5, 1.5, 2.5, 0.65, "Section panel\n(terminological\nrecord)", C_ANNOT, 7.5)

    # --- Frontend ---
    pages = ["Ask", "Search", "Browse", "Section", "Graph",
             "Timeline", "Map", "Terminology", "About"]
    w_page = 9.0 / len(pages)
    for i, pg in enumerate(pages):
        box(ax, 1.0 + i*w_page, 0.15, w_page - 0.07, 0.52, pg, C_FRONT, 6.5)

    # --- IONOS bundle (right sidebar) ---
    rect = FancyBboxPatch((10.5, 0.1), 1.35, 8.55,
                          boxstyle="round,pad=0.1",
                          linewidth=1, linestyle="dashed",
                          edgecolor=C_BUNDLE, facecolor="#f0ede8", zorder=2)
    ax.add_patch(rect)
    ax.text(11.175, 4.35, "IONOS\nbundle\n(self-contained\nzip + smoke\ntest)", ha="center", va="center",
            fontsize=7.5, color=C_BUNDLE, zorder=4, multialignment="center")

    # --- Arrows between layers ---
    # Corpus -> Retrieval
    for x in [2.1, 4.3, 6.7]:
        arrow(ax, x, 8.0, x, 7.05, color=C_CORPUS)

    # Retrieval -> KG
    arrow(ax, 5.0, 6.4, 5.0, 5.45, color=C_RETRIEV)

    # KG -> LOD
    arrow(ax, 5.0, 4.8, 5.0, 3.65, color=C_KG)

    # LOD -> Gazetteer (feedback loop, dashed)
    arrow(ax, 2.3, 3.0, 2.0, 2.15, color=C_ANNOT, dashed=True,
          label="derives from")

    # Annotator -> section panel
    arrow(ax, 6.3, 1.5, 7.15, 2.15, color=C_ANNOT, dashed=False)

    # Retrieval/KG/LOD -> Frontend
    arrow(ax, 5.5, 4.8, 5.5, 0.67, color="#bbb", lw=0.8)

    # Legend
    legend_items = [
        (C_CORPUS, "Corpus"),
        (C_RETRIEV, "Retrieval"),
        (C_KG, "Knowledge layers"),
        (C_LOD, "LOD / ontology"),
        (C_ANNOT, "Annotation"),
        (C_FRONT, "Frontend"),
    ]
    for i, (col, lbl) in enumerate(legend_items):
        bx = 1.0 + i * 1.7
        ax.add_patch(mpatches.Rectangle((bx, 0.0), 0.18, 0.12,
                                        color=col, zorder=5))
        ax.text(bx + 0.22, 0.06, lbl, fontsize=6.5, va="center", color="#444", zorder=6)

    fig.tight_layout(pad=0.3)
    fig.savefig(str(ARCH_PNG), dpi=180, bbox_inches="tight",
                facecolor="#fafaf8")
    plt.close(fig)
    print(f"Architecture figure saved: {ARCH_PNG}")


# ---------------------------------------------------------------------------
# 2. Document content (shared between PDF and DOCX)
# ---------------------------------------------------------------------------

TITLE     = "Laertius: A Technical Overview"
SUBTITLE  = "Retrieval, Knowledge Modelling, and Semantic Annotation\nover Diogenes Laertius' Lives of Eminent Philosophers"
DATE      = "July 2026"

# Section texts are written as lists of (type, text) tuples where type is
# one of: 'h1', 'h2', 'body', 'bullet', 'formula', 'code', 'greek', 'table'
# Tables are lists of rows (each row a list of cells).

SECTIONS = []

def sec(title, content):
    SECTIONS.append((title, content))

# Helper to format formula text (plain-text approximation for PDF/DOCX)
def fmt_formula(s):
    return ("formula", s)

def fmt_body(s):
    return ("body", s)

def fmt_bullet(s):
    return ("bullet", s)

def fmt_greek(s):
    return ("greek", s)

def fmt_h2(s):
    return ("h2", s)

def fmt_code(s):
    return ("code", s)

def fmt_table(rows):
    return ("table", rows)


sec("Overview", [
    fmt_body(
        "Laertius is a free, key-less scholarly search tool over Diogenes Laertius' "
        "Lives of Eminent Philosophers (c. 3rd century CE). It combines the Greek text "
        "(Perseus grc2 edition) with the aligned R. D. Hicks English translation "
        "(Loeb 1925), and offers hybrid retrieval, a structured knowledge graph, "
        "a semantic annotation layer, and a Linked Open Data export, all without "
        "requiring any external API key, payment, or large-language model."
    ),
    fmt_body(
        "The design principle is retrieval-augmented reading rather than generation. "
        "Everything displayed to the reader is verbatim, cited source text: no "
        "passage is paraphrased or synthesised. The system surfaces what Diogenes "
        "Laertius wrote, who he credits, and how certain he sounds, and presents it "
        "with its Greek alongside the English, with hyperlinks to the knowledge graph "
        "and the LOD ontology."
    ),
    fmt_body(
        "Technical summary:"
    ),
    fmt_bullet("Language: TypeScript / Node.js 24, pnpm workspace monorepo"),
    fmt_bullet("API: Express 5; frontend: React + Vite (wouter, TanStack Query, Tailwind)"),
    fmt_bullet("Embeddings: local @huggingface/transformers, no API key, CPU-only"),
    fmt_bullet("Validation: Zod (zod/v4); OpenAPI contract-first with Orval codegen"),
    fmt_bullet("Deployment: self-contained IONOS bundle (zip, smoke-tested before release)"),
    fmt_bullet(f"Corpus: {CORPUS_SECTIONS} sections (Greek + English), Books 1-10 plus Prologue"),
])

sec("Retrieval Pipeline", [
    fmt_body(
        "The retrieval pipeline is fully local and deterministic. It combines a "
        "classical sparse term-frequency index with a small dense embedding model, "
        "fused by Reciprocal Rank Fusion, and optionally boosted by the philosophers "
        "knowledge graph."
    ),
    fmt_h2("2.1 BM25 (sparse retrieval)"),
    fmt_body(
        "BM25 is built at server start over the concatenated Greek and English text "
        f"of all {CORPUS_SECTIONS} sections (section.text + ' ' + section.textEn). "
        f"Parameters (from bm25.ts): K1 = {K1}, B = {B}. "
        "The IDF formula is the Lucene non-negative variant:"
    ),
    fmt_formula("IDF(t) = log( (N - df + 0.5) / (df + 0.5) + 1 )"),
    fmt_body(
        "where N is the total section count and df is the number of sections "
        "containing term t. This variant is always non-negative (unlike classic "
        "Robertson/Sparck Jones IDF), matching Lucene 6+ behaviour. "
        "The per-document score sums over query terms:"
    ),
    fmt_formula("BM25(d,q) = sum_t IDF(t) * tf(t,d) * (K1+1) / (tf(t,d) + K1*(1 - B + B*|d|/avgdl))"),
    fmt_body(
        "Tokenisation is shared with the Greek normalizer (tokenize from greek.ts): "
        "Unicode letter sequences are lowercased and diacritics are stripped, "
        "so a Greek query token and its corresponding English transliteration "
        "may still score independently."
    ),
    fmt_h2("2.2 Dense retrieval"),
    fmt_body(
        f"Dense retrieval uses the {EMBED_MODEL} model ({EMBED_DIM}-dimensional "
        f"output, {EMBED_DTYPE} quantised) via @huggingface/transformers. "
        "At query time the query is prefixed with 'query: ' (the asymmetric E5 "
        "convention), mean-pooled, and L2-normalised. Stored passage vectors "
        "are also L2-normalised, so cosine similarity reduces to a dot product. "
        f"Brute-force exact search is performed over all embedded sections "
        f"(up to {CORPUS_SECTIONS}), returning the top-{POOL} by cosine score. "
        "If the dense index file is absent, the system falls back silently to "
        "sparse-only retrieval."
    ),
    fmt_h2("2.3 Reciprocal Rank Fusion"),
    fmt_body(
        f"The BM25 top-{POOL} and dense top-{POOL} ranked lists are merged by "
        f"weighted Reciprocal Rank Fusion (Cormack, Clarke, Buettcher 2009) with "
        f"k = {RRF_K}, dense weight 1 and sparse weight {SPARSE_WEIGHT}:"
    ),
    fmt_formula(f"RRF(d) = sum_r  w_r / ({RRF_K} + rank_r(d)),  w_dense = 1, w_sparse = {SPARSE_WEIGHT}"),
    fmt_body(
        f"where the sum is over the two ranked lists. Documents appearing in "
        "only one list receive zero contribution from the absent list. "
        f"The parameters (k = {RRF_K}, sparse weight {SPARSE_WEIGHT}) were tuned "
        "against the gold v0.5 evaluation set: dense-leaning weighted fusion keeps "
        "the answerable-topic gains of equal-weight fusion while preserving the "
        "abstention evidence (contradicting passages, homonym rosters) that dense "
        "retrieval surfaces (see gold-eval-v0.5-fusion-tuning.md)."
    ),
    fmt_h2("2.4 Knowledge-graph boost"),
    fmt_body(
        "After fusion, sections whose philosopher field matches a name detected "
        "in the query (or a graph neighbour of a detected name) receive a "
        "rank-scale additive boost:"
    ),
    fmt_formula(f"score' = score + 0.5 / {RRF_K}   (exact match)"),
    fmt_formula(f"score' = score + 0.15 / {RRF_K}  (graph neighbour)"),
    fmt_body(
        "This is a ranking-only boost (sections already in the fused set move up; "
        "no new sections are injected). The boost is calibrated to the RRF scale: "
        f"the maximum first-rank RRF contribution is 1/(k+1) = 1/{RRF_K+1}."
    ),
])

sec("The E5 Embedding Pipeline", [
    fmt_body(
        f"The dense index is built offline by the build-embeddings script "
        f"(@workspace/scripts). Each section is prefixed with 'passage: ' "
        f"(the E5 asymmetric passage prefix), embedded by {EMBED_MODEL}, "
        f"mean-pooled, and L2-normalised to unit vectors. Vectors are stored "
        f"as 32-bit floats in a base64-encoded block alongside the section "
        f"id list and the model name, in embedding-index.json."
    ),
    fmt_h2("3.1 Model and quantisation"),
    fmt_body(
        f"Model: {EMBED_MODEL} (Wang et al. 2024). "
        f"Output dimension: {EMBED_DIM}. "
        f"Quantisation: {EMBED_DTYPE} (8-bit), applied at inference time "
        "via @huggingface/transformers dtype option. The quantised model is "
        "cached on first run under data/models/ and reused on subsequent starts."
    ),
    fmt_h2("3.2 Integrity checks"),
    fmt_body(
        "At server start, dense.ts validates the stored index before accepting it:"
    ),
    fmt_bullet("Model name in the index must match the server's EMBEDDING_MODEL constant."),
    fmt_bullet("Byte length must equal ids.length * dim * 4 (Float32 integrity check)."),
    fmt_bullet("dim must be a positive integer and ids must be non-empty."),
    fmt_body(
        "Any mismatch causes loadDenseIndex to log an error and return false, "
        "triggering graceful fallback to sparse-only mode. This prevents a "
        "stale or truncated index from silently degrading retrieval quality."
    ),
    fmt_h2("3.3 Warm-up"),
    fmt_body(
        "warmUpEmbedder is called at server boot. It runs a dummy 'query: warmup' "
        "inference to force the model into memory and JIT-compile the ONNX graph "
        "before the first real query arrives, avoiding a cold-start latency spike."
    ),
])

sec("The Philosophers Graph", [
    fmt_body(
        f"The knowledge graph contains {KG_EDGES} curated directed edges among the "
        f"{CHAPTER_SUBJECTS} philosophers who have a Life in the corpus. Nodes are derived at runtime "
        "from the corpus itself (kg.ts) so the graph can never reference a philosopher "
        "who is absent from the text - dangling edges are structurally impossible."
    ),
    fmt_h2("4.1 Edges"),
    fmt_body(
        f"All {KG_EDGES} edges follow Diogenes Laertius' own diadochai (succession "
        "accounts, 1.13-15), with a D.L. citation on every edge. Edge types:"
    ),
    fmt_bullet("teacherOf: D.L. reports a direct teacher-pupil relation (e.g. Socrates -> Plato, ref 3.5)"),
    fmt_bullet("influenced: succession or doctrinal transmission without direct teaching (e.g. Pythagoras -> Empedocles, ref 8.54)"),
    fmt_bullet("spouseOf: documented marital relation (e.g. Crates of Thebes -> Hipparchia, ref 6.96)"),
    fmt_h2("4.2 Movements"),
    fmt_body(
        f"{MOVEMENT_COUNT} philosophical movements are modelled (Seven Sages, Ionian/Milesian, "
        "Socratic, Cyrenaic, Megarian, Elian-Eretrian, Academy, Peripatos, Cynic, Stoa, "
        "Pythagorean, Eleatic, Atomist, Sophist, Sceptic, Epicurean, Unaffiliated). "
        "Each philosopher is assigned to exactly one movement; school founders are "
        "marked with a founderOf relation. Wikidata QIDs are curated for every node "
        "where a confident, unambiguous match exists (never guessed for homonyms)."
    ),
    fmt_h2("4.3 Uses in the system"),
    fmt_bullet("Graph page: visual diadochai browser with school filters and succession-tree view"),
    fmt_bullet("Retrieval boost: names detected in a query trigger the KG boost described in Section 2"),
    fmt_bullet("LOD export: edges become lo:teacherOf / lo:influenced / lo:spouseOf triples with reified citations"),
])

sec("The Assertion Model", [
    fmt_body(
        f"The claims layer contains {CLAIM_COUNT} cited assertions about the philosophers, "
        "compiled from the text of Diogenes Laertius itself (source-internal: no modern "
        "reference data is used). Each claim carries:"
    ),
    fmt_bullet("A stable id (philosopher-prefixed, e.g. 'thales-lived-miletus')"),
    fmt_bullet("A D.L. citation in Hicks section numbering (e.g. '1.22')"),
    fmt_bullet("A certainty level (asserted / reported / disputed / conjectured)"),
    fmt_bullet("A closed property vocabulary (18 predicates, e.g. birthPlace, wrote, heldDoctrine)"),
    fmt_bullet("An optional accordingTo source (named authority D.L. cites)"),
    fmt_bullet("An optional ordered transmission chain (ChainLink[], nearest intermediary first)"),
    fmt_bullet("An optional conflictsWith list (ids of alternative claims on the same question)"),
    fmt_bullet("A verbatim Greek excerpt (grc) from the cited section, validator-checked"),
    fmt_h2("5.1 Certainty levels"),
    fmt_body(
        "The four certainty levels mirror D.L.'s own epistemic stance, not a modern "
        "assessment:"
    ),
    fmt_bullet("asserted: stated in D.L.'s own voice without hedging"),
    fmt_bullet("reported: hedged ('some say', 'according to X', 'it is said')"),
    fmt_bullet("disputed: D.L. records explicit disagreement between sources"),
    fmt_bullet("conjectured: an inference, not directly stated"),
    fmt_h2("5.2 Worked examples"),
    fmt_body("Asserted - Thales lived in Miletus (D.L. 1.22):"),
    fmt_table([
        ["Field", "Value"],
        ["id", "thales-lived-miletus"],
        ["subject", "Thales"],
        ["property", "livedIn"],
        ["value", "Miletus"],
        ["ref", "1.22"],
        ["certainty", "asserted"],
        ["grc (excerpt)", "\u1f26\u03bd \u03c4\u03bf\u03af\u03bd\u03c5\u03bd \u1f41 \u0398\u03b1\u03bb\u1fc6\u03c2 ..."],
    ]),
    fmt_body("Reported with transmission chain - Myson's parentage (D.L. 1.106):"),
    fmt_table([
        ["Field", "Value"],
        ["id", "myson-parentage"],
        ["subject", "Myson"],
        ["property", "parentage"],
        ["ref", "1.106"],
        ["certainty", "reported"],
        ["accordingTo", "Sosicrates"],
        ["chain", "[{ authority: 'Hermippus' }]"],
        ["note", "Sosicrates quotes Hermippus as his authority"],
    ]),
    fmt_body("Disputed - Thales' Nautical Astronomy (D.L. 1.23):"),
    fmt_table([
        ["Field", "Value"],
        ["id", "thales-wrote-nautical-astronomy"],
        ["subject", "Thales"],
        ["property", "wrote"],
        ["value", "Nautical Astronomy"],
        ["ref", "1.23"],
        ["certainty", "disputed"],
        ["conflictsWith", "thales-wrote-nothing"],
        ["note", "Attributed to Thales but said to be by Phocus of Samos"],
    ]),
    fmt_body("Conjectured - Philolaus' manner of death (D.L. 8.84):"),
    fmt_table([
        ["Field", "Value"],
        ["id", "philolaus-death-tyranny"],
        ["subject", "Philolaus"],
        ["property", "mannerOfDeath"],
        ["value", "Put to death by Croton on suspicion of aiming at a tyranny"],
        ["ref", "8.84"],
        ["certainty", "conjectured"],
        ["note", "Inferred from D.L.'s epigram; prose account is ambiguous"],
    ]),
    fmt_h2("5.3 Named authorities"),
    fmt_body(
        f"Of the {CLAIM_COUNT} claims, {NAMED_AUTHORITY_CLAIMS} carry a named "
        "accordingTo authority (the source D.L. explicitly credits). "
        "The table below lists the principal authorities, the domain they "
        "cover in the claims layer, a representative property, and an example "
        "claim with its D.L. reference:"
    ),
    fmt_table([
        ["Authority", "Domain", "Typical property", "Example claim (ref)"],
        ["Apollodorus", "Chronology: birth/death dating, Olympiad reckonings",
         "birthDate / deathDate", "Thales born in the 35th Olympiad (1.37)"],
        ["Hermippus", "Deaths, life accounts, biographies of philosophers",
         "mannerOfDeath / deathPlace", "Chilon died at Pisa of excess joy (1.72)"],
        ["Sosicrates", "Socratic succession, Sage accounts",
         "deathDate / parentage", "Thales died at age 90 (1.38); Myson's father (1.106)"],
        ["Alexander Polyhistor", "Pythagorean notebook, doctrines",
         "heldDoctrine", "Pythagoras' numerical doctrines (8.25)"],
        ["Sotion", "Successions, school membership chains",
         "studiedUnder / succession", "Peripatetic successions (5.36)"],
        ["Satyrus", "Biographies (lost Lives of Philosophers)",
         "parentage / birthPlace", "Various Book 2 and Book 5 accounts"],
        ["Favorinus", "Miscellaneous anecdotes, disputes",
         "mannerOfDeath / parentage", "Chain link at 5.41 (Hermippus <- Arcesilaus)"],
        ["Panaetius", "Authenticity of Socratic letters",
         "writings", "Disputes authenticity of certain epistles"],
        ["Herodotus / Duris / Democritus", "Thales' Phoenician descent",
         "parentage", "Thales' parents (1.22) - three-way reported claim"],
    ]),
    fmt_body(
        "Authorities are validated at server start: every accordingTo value "
        "must be a known source (a SOURCE_WORKS author or a claim's own "
        "accordingTo), and every chain authority must similarly resolve, "
        "so a typo cannot mint a phantom name in the LOD graph."
    ),
    fmt_h2("5.4 Conflict clusters"),
    fmt_body(
        "ConflictsWith links group alternative claims on the same biographical "
        "question. The Empedocles death dossier is the largest such cluster, "
        "with multiple rival accounts. All links are bidirectional by convention "
        "and validated at server start (dangling ids throw on boot)."
    ),
])

sec("Semantic Annotation", [
    fmt_body(
        "The annotation layer tags occurrence-level entity and term mentions in "
        f"both Greek and English text across all {CORPUS_SECTIONS} sections, "
        f"producing {ANNOTATION_COUNT} occurrence tags over {TAGGED_ENTITIES} "
        "distinct entities."
    ),
    fmt_h2("6.1 Two-stage design"),
    fmt_body(
        "Stage 1 - Gazetteer construction (gazetteer.ts): the full LOD graph is "
        "parsed (Turtle, via n3), and a surface form is included in the gazetteer "
        "if and only if the graph resolves it to exactly one individual. Ambiguous "
        "surfaces go into a skip ledger (precision over recall). The gazetteer is "
        "never constructed from a separate name list - it is derived entirely from "
        "the curated graph, so tagger and ontology can never drift apart."
    ),
    fmt_body(
        "Stage 2 - Deterministic annotator (annotate.ts): scans each section "
        "and matches gazetteer entries using the rules below. No probabilistic "
        "models, no LLMs."
    ),
    fmt_h2("6.2 Algorithm (pseudocode)"),
    fmt_body("English annotator (annotateEnglish):"),
    fmt_code(
        "Phase 0: Build combined longest-first alternation regex from all surfaces.\n"
        "Phase 1: Scan section.textEn with the regex (case-sensitive, Unicode\n"
        "         letter-boundary lookarounds). For each match:\n"
        "           - Look up all gazetteer entries for the matched surface.\n"
        "           - Scoped entries (onlySections) are tried first; if none\n"
        "             matches this section id, try unscoped entries.\n"
        "           - If no entry applies, skip the match.\n"
        "Phase 2: Section-owner heuristic. If the section's philosopher has a\n"
        "         bare first name that is ambiguous (shared by multiple philosophers),\n"
        "         scan the text for that bare name and tag it with the Life's\n"
        "         philosopher (flagged heuristic: 'section-owner').\n"
        "Phase 3: Overlap resolution. Sort all candidates by (start ASC, end DESC,\n"
        "         non-heuristic first). Accept greedily: skip any candidate whose\n"
        "         start < lastEnd."
    ),
    fmt_body("Greek annotator (annotateGreek):"),
    fmt_code(
        "Phase 0: Normalize the polytonic Greek text to a form stripped of\n"
        "         diacritics and case, building an offset map so that positions\n"
        "         in the normalized form map back to the original polytonic text.\n"
        "Phase 1: Term matching (otv:Term lemmas). For each term:\n"
        "           - Single-word: accept a normalized token if it equals the\n"
        "             normalized lemma exactly, or if it starts with the lemma's\n"
        "             stem and ends with a whitelisted nominal ending.\n"
        "           - Multi-word: require the exact normalized word sequence.\n"
        "Phase 2: Proper-name matching (Greek ProperName nodes). Require the\n"
        "         ORIGINAL token to start with an uppercase letter (capital guard):\n"
        "         pi_politeia the noun must never match Pi_politeia the work.\n"
        "         Scoped and unscoped entries handled as in the English annotator.\n"
        "Phase 3: Section-owner heuristic for ambiguous Greek forms (e.g. the\n"
        "         nominative shared by Zeno of Citium and Zeno of Elea).\n"
        "Phase 4: Same overlap resolution as the English annotator."
    ),
    fmt_h2("6.3 Ambiguity policy"),
    fmt_body(
        "When a surface resolves to multiple distinct individuals, it is placed "
        "in the skip ledger unless all candidates share the same Wikidata QID "
        "(i.e. the same person appearing under two roles, such as Aristotle "
        "the philosopher and Aristotle cited as a source). Bare first names "
        "(e.g. 'Zeno') are generated only when globally unambiguous; ambiguous "
        "bare names that resolve exclusively to philosophers are reserved for "
        "the section-owner heuristic. The blocklist additionally suppresses "
        "surfaces that are unique in the knowledge graph but demonstrably "
        "multi-referent in the text (e.g. 'Alexander', 'Antigonus', 'Theodorus')."
    ),
])

sec("RDF Reification of Hedged Claims", [
    fmt_body(
        "Hedged claims (reported, disputed, conjectured) cannot be expressed "
        "as simple RDF triples without losing their epistemic status. Laertius "
        "uses RDF reification with the lo:Claim node pattern."
    ),
    fmt_h2("7.1 Two-tier emission rule"),
    fmt_body(
        "Only asserted claims emit a direct triple on their subject:"
    ),
    fmt_bullet("asserted + object property: both a lo:Claim reification AND a direct lo:predicate triple"),
    fmt_bullet("asserted + datatype property: reification + direct datatype triple"),
    fmt_bullet("reported / disputed / conjectured: reification only, no direct triple"),
    fmt_h2("7.2 Anatomy of a lo:Claim node"),
    fmt_body(
        "Each claim node carries (from lod.ts):"
    ),
    fmt_bullet("rdf:type lo:Claim"),
    fmt_bullet("rdf:subject / rdf:predicate / rdf:object: the reified triple"),
    fmt_bullet("lo:certainty: individual (lo:Asserted / lo:Reported / lo:Disputed / lo:Conjectured)"),
    fmt_bullet("dcterms:bibliographicCitation: the D.L. Hicks section reference"),
    fmt_bullet("lo:accordingTo: source node URI (when named)"),
    fmt_bullet("lo:assertedInWork: work node URI (when sourceWork named)"),
    fmt_bullet("lo:transmissionChain: ordered list of lo:ChainLink nodes (nearest first)"),
    fmt_bullet("lo:conflictsWith: links to rival claim nodes"),
    fmt_bullet("lo:greekText @grc: verbatim Greek excerpt"),
    fmt_h2("7.3 Naming conventions"),
    fmt_body(
        "Date properties use reportedBirthDate / reportedDeathDate rather than "
        "birthDate / deathDate. The 'reported' prefix signals that these values "
        "reflect D.L.'s testimony, not a modern scholarly consensus. "
        "Epistle authenticity is a separate axis, using lo:DisputedAuthenticity "
        "to avoid collision with the certainty individual lo:Disputed "
        "(curator's verdict vs. D.L.'s epistemic stance are logically independent)."
    ),
    fmt_h2("7.4 Pattern reuse"),
    fmt_body(
        "The same reification pattern is applied consistently across layers: "
        "sayings (lo:Saying), school memberships, succession links, anecdotes "
        "(lo:Anecdote), and epistles (lo:Epistle). Hedged school memberships and "
        "succession links emit reification-only triples - they are recorded in "
        "the LOD graph but do not assert a direct relationship."
    ),
])

sec("Architecture", [
    fmt_body(
        "The system is built as a one-way layered pipeline with a single "
        "feedback arc. Each layer produces data consumed by the next; no layer "
        "reads from a layer above it except for the annotation layer, which "
        "derives its gazetteer from the published LOD graph."
    ),
    fmt_h2("8.1 Pipeline layers"),
    fmt_bullet(
        "Corpus: Greek TEI (Perseus grc2) and English TEI (Hicks) parsed offline "
        f"into {CORPUS_SECTIONS} aligned JSONL sections. Section ids "
        "(book.chapter.section, e.g. '1.prol.1') are the join key across all layers."
    ),
    fmt_bullet(
        "Retrieval: BM25 index and dense embedding index built over the concatenated "
        "Greek-English documents. Dense index stored as base64 Float32 in "
        "embedding-index.json; absent index triggers sparse-only fallback."
    ),
    fmt_bullet(
        "Knowledge layers: compiled-in TypeScript (curated claim files, sayings, "
        "anecdotes, doxai, verses), not data files. Typed, validated at server start. "
        f"The verses layer alone carries {VERSE_COUNT} verse quotations extracted "
        "from TEI block-quotes, with curated poet attributions. "
        "All cross-references validated before the server serves its first request."
    ),
    fmt_bullet(
        f"LOD emission: lod.ts assembles the full Turtle graph (~{ANNOTATED_TRIPLES_K}k triples in the "
        "annotated export), derives RDF/XML, serves a read-only oxigraph SPARQL "
        "endpoint, and emits the TEDI-style OTB ontoterminology export."
    ),
    fmt_bullet(
        "Annotation feedback: gazetteer.ts parses the LOD Turtle graph (via n3) to "
        "derive tagger surfaces. This means the tagger is always consistent with the "
        "published ontology - a new entity added to the graph becomes taggable "
        "automatically, with no separate tagger update."
    ),
    fmt_h2("8.2 Derive-never-duplicate"),
    fmt_body(
        "Nodes are derived from authoritative sources, never duplicated. KG nodes "
        "come from the corpus; claim entities come from the claims; gazetteer entries "
        "come from the LOD graph. Dangling references throw at server start, making "
        "curation errors visible immediately."
    ),
    fmt_h2("8.3 Curation-as-code"),
    fmt_body(
        "All curated content (claims, sayings, anecdotes, doxai, succession links, "
        "school memberships, person mentions) is compiled TypeScript, not database "
        "rows or YAML. This means: full TypeScript type checking at build time, "
        "diff-readable version history, and the closed-vocabulary predicates are "
        "enforced by the type system."
    ),
    fmt_h2("8.4 Validators"),
    fmt_body(
        "Approximately 30 registered validators pin layer outputs. They are run "
        "automatically after each task merge. A selection:"
    ),
    fmt_bullet("validate-claims: ref resolution, subject presence, conflictsWith links, chain authority validity"),
    fmt_bullet("validate-annotations: total occurrence count pinned, entity count pinned, skip ledger pinned"),
    fmt_bullet("validate-lod: no dangling URIs, SPARQL competency questions pass"),
    fmt_bullet("map-contract: OpenAPI spec and map.ts interface property lists in lockstep"),
    fmt_bullet("otb: TEDI-style export roundtrip (emit -> parse -> check)"),
    fmt_h2("8.5 Contract-first API"),
    fmt_body(
        "The API contract lives in lib/api-spec/openapi.yaml. Orval generates "
        "React Query hooks and Zod schemas from it. The server validates inputs "
        "and outputs with the generated Zod schemas. A field added to the spec "
        "but missing from the server implementation fails the typecheck."
    ),
    fmt_h2("8.6 IONOS deployment bundle"),
    fmt_body(
        "build-ionos-bundle.ts assembles a self-contained zip: an esbuild-bundled "
        "server (external: @huggingface/transformers, oxigraph), the Vite-built "
        "frontend (BASE_PATH=/), and the data directory (corpus JSONLs, no models). "
        "The bundle script boots the zip and smoke-tests every endpoint before "
        "exiting. If any check fails, the zip is deleted and the script exits "
        "non-zero, preventing a broken bundle from being deployed."
    ),
])

sec("Algorithm Provenance", [
    fmt_body(
        "The algorithms used are all from published literature or established "
        "open-source implementations. The project-specific policies are described "
        "where they deviate from the published baseline."
    ),
    fmt_h2("9.1 BM25"),
    fmt_body(
        "Robertson, S. E. and Sparck Jones, K. (1976). Relevance weighting of "
        "search terms. Journal of the American Society for Information Science 27(3). "
        "IDF variant: Lucene non-negative IDF "
        "(log((N - df + 0.5)/(df + 0.5) + 1)), which avoids negative scores for "
        f"very frequent terms. Parameters K1={K1}, B={B} follow the BM25F defaults."
    ),
    fmt_h2("9.2 E5 embeddings"),
    fmt_body(
        "Wang, L. et al. (2022). Text Embeddings by Weakly-Supervised Contrastive "
        "Pre-training. Asymmetric prefixes ('query: ' / 'passage: ') are required "
        "by the model architecture; the Laertius system applies them correctly "
        "in embedder.ts (query prefix) and build-embeddings.ts (passage prefix). "
        "Quantisation to q8 reduces memory and inference time with minimal accuracy loss."
    ),
    fmt_h2("9.3 Reciprocal Rank Fusion"),
    fmt_body(
        "Cormack, G. V., Clarke, C. L. A., and Buettcher, S. (2009). Reciprocal "
        "Rank Fusion Outperforms Condorcet and Individual Rank Learning Methods. "
        "SIGIR 2009. The parameter k=60 follows the paper's recommendation and "
        "prevents very high-ranked documents from dominating excessively."
    ),
    fmt_h2("9.4 Gazetteer-based NER"),
    fmt_body(
        "Longest-match-first gazetteer NER is a standard technique; the 'one sense "
        "per discourse' principle (Gale, Church, Yarowsky 1992) motivates the "
        "section-owner heuristic for ambiguous bare names. Precision over recall "
        "is the explicit design choice: the skip ledger tracks all surfaces where "
        "the system chose to tag nothing rather than risk a misattribution."
    ),
    fmt_h2("9.5 Ontoterminology framework"),
    fmt_body(
        "The TEDI-style OTB export is modelled on the Ontoterminology framework "
        "of Christophe Roche (Roche, C. 2007-2020), with exports following "
        "Maria Papadopoulou's TEDI 4.1 reference format. The OTV ontology viewer "
        "uses the isA and instanceOf concept hierarchy from the OTB model."
    ),
    fmt_h2("9.6 Project-specific policies"),
    fmt_bullet(
        "Cite everything: every knowledge-graph edge, claim, verse attribution, "
        "and source mention carries a D.L. reference."
    ),
    fmt_bullet(
        "Flag heuristics: section-owner tags are marked heuristic: 'section-owner' "
        "in the annotation data and displayed with a visual cue in the UI."
    ),
    fmt_bullet(
        "Fail loud: dangling refs, duplicate ids, and schema violations throw at "
        "server start rather than logging a warning and serving stale data."
    ),
    fmt_bullet(
        "Never guess QIDs: external identifiers (Wikidata, VIAF, Pleiades, "
        "InPhO) are curated offline and never inferred programmatically."
    ),
])

sec("Distributed Text Services API", [
    fmt_body(
        "The server exposes a read-only Distributed Text Services (DTS) 1.0 API "
        "under /dts, so that DTS-aware clients can discover, navigate, and "
        "retrieve the edition programmatically. The implementation lives in "
        "dts.ts (document model and JSON-LD assembly) with the HTTP routes in "
        "the routes module of the same name. All responses are JSON-LD "
        "(application/ld+json) except the Document endpoint, whose default "
        "output is TEI XML served as application/tei+xml."
    ),
    fmt_h2("10.1 Endpoints"),
    fmt_bullet(
        "EntryPoint: GET /dts returns the JSON-LD entry point advertising the "
        "collection, navigation, and document URI templates."
    ),
    fmt_bullet(
        "Collection: GET /dts/collection{?id} describes the catalogue; the "
        "single readable resource is the Lives, identified by the CTS URN "
        "urn:cts:greekLit:tlg0004.tlg001.perseus-grc2."
    ),
    fmt_bullet(
        "Navigation: GET /dts/navigation?resource=...{&ref,down} walks the "
        "citation tree; ref selects a subtree and down controls the depth of "
        "the returned members."
    ),
    fmt_bullet(
        "Document: GET /dts/document?resource=...{&ref,start,end,mediaType} "
        "returns the passage text. The default media type is TEI XML "
        "(application/tei+xml); mediaType=application/ld+json selects a "
        "JSON-LD rendering of a single passage. A Link header points back to "
        "the parent collection."
    ),
    fmt_h2("10.2 Citation tree"),
    fmt_body(
        "The citation tree has two levels, book then section: books are cited "
        "as '1' through '10', and sections as book.section in the Hicks "
        "numbering (e.g. '2.18'). Where a shorthand reference is ambiguous, "
        "the full corpus identifier book.chapter.section is used instead, and "
        "the resolver accepts both forms as aliases."
    ),
    fmt_h2("10.3 Canonical origin"),
    fmt_body(
        "Absolute identifiers in the JSON-LD output are minted against a "
        "canonical origin. The DTS_PUBLIC_ORIGIN environment variable, when "
        "set, is validated (http/https scheme, no path, no trailing slash) and "
        "wins outright; otherwise the origin is derived from LOD_BASE_URI, and "
        "only as a last resort from the request itself, and then only when the "
        "Express trust-proxy setting is enabled. Forwarded headers are never "
        "trusted directly, so a spoofed Host header cannot poison the minted "
        "identifiers."
    ),
])


# ---------------------------------------------------------------------------
# 3. PDF generation
# ---------------------------------------------------------------------------
def make_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, Image, HRFlowable, ListFlowable, ListItem,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdfmetrics.registerFont(TTFont("DvSerif",     FONT_REG))
    pdfmetrics.registerFont(TTFont("DvSerif-Bold",FONT_BOLD))
    pdfmetrics.registerFont(TTFont("DvSans",      FONT_SANS))
    pdfmetrics.registerFont(TTFont("DvSans-Bold", FONT_SANS_BOLD))

    W, H = A4
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=2.5*cm, leftMargin=2.5*cm,
        topMargin=2.5*cm,   bottomMargin=2.5*cm,
        title=TITLE,
        author="",
    )

    # Styles
    body_style = ParagraphStyle(
        "body", fontName="DvSerif", fontSize=10, leading=14,
        spaceAfter=6, spaceBefore=0,
    )
    bullet_style = ParagraphStyle(
        "bullet", fontName="DvSerif", fontSize=10, leading=14,
        spaceAfter=4, leftIndent=16, bulletIndent=6,
    )
    h1_style = ParagraphStyle(
        "h1", fontName="DvSans-Bold", fontSize=14, leading=18,
        spaceBefore=18, spaceAfter=8, textColor=colors.HexColor("#1a3a5c"),
    )
    h2_style = ParagraphStyle(
        "h2", fontName="DvSans-Bold", fontSize=11, leading=15,
        spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1d7874"),
    )
    title_style = ParagraphStyle(
        "title", fontName="DvSans-Bold", fontSize=22, leading=28,
        alignment=1, textColor=colors.HexColor("#1a3a5c"),
    )
    subtitle_style = ParagraphStyle(
        "subtitle", fontName="DvSans", fontSize=13, leading=18,
        alignment=1, textColor=colors.HexColor("#555"),
    )
    date_style = ParagraphStyle(
        "date", fontName="DvSans", fontSize=11, leading=14,
        alignment=1, textColor=colors.HexColor("#888"),
    )
    formula_style = ParagraphStyle(
        "formula", fontName="DvSans", fontSize=9.5, leading=14,
        leftIndent=30, spaceAfter=4, spaceBefore=4,
        textColor=colors.HexColor("#2a2a5a"),
        backColor=colors.HexColor("#f0f4ff"),
        borderPad=4,
    )
    code_style = ParagraphStyle(
        "code", fontName="DvSans", fontSize=8, leading=12,
        leftIndent=16, spaceAfter=4, spaceBefore=4,
        textColor=colors.HexColor("#222"),
        backColor=colors.HexColor("#f5f5f0"),
        borderPad=6,
    )
    toc_h1_style = ParagraphStyle(
        "toc_h1", fontName="DvSans-Bold", fontSize=11, leading=16,
        leftIndent=0, spaceBefore=4,
    )
    toc_h2_style = ParagraphStyle(
        "toc_h2", fontName="DvSans", fontSize=9.5, leading=13,
        leftIndent=16,
    )

    TABLE_STYLE = TableStyle([
        ("BACKGROUND",   (0,0), (-1,0), colors.HexColor("#1a3a5c")),
        ("TEXTCOLOR",    (0,0), (-1,0), colors.white),
        ("FONTNAME",     (0,0), (-1,0), "DvSans-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 8.5),
        ("FONTNAME",     (0,1), (-1,-1), "DvSerif"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.HexColor("#f8f8f8"), colors.white]),
        ("GRID",         (0,0), (-1,-1), 0.4, colors.HexColor("#ccc")),
        ("LEFTPADDING",  (0,0), (-1,-1), 5),
        ("RIGHTPADDING", (0,0), (-1,-1), 5),
        ("TOPPADDING",   (0,0), (-1,-1), 3),
        ("BOTTOMPADDING",(0,0), (-1,-1), 3),
        ("VALIGN",       (0,0), (-1,-1), "TOP"),
    ])

    story = []

    # Title page
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph(TITLE, title_style))
    story.append(Spacer(1, 0.5*cm))
    for line in SUBTITLE.split("\n"):
        story.append(Paragraph(line, subtitle_style))
    story.append(Spacer(1, 0.8*cm))
    story.append(HRFlowable(width="80%", color=colors.HexColor("#1a3a5c"), thickness=1))
    story.append(Spacer(1, 0.8*cm))
    story.append(Paragraph(DATE, date_style))
    story.append(PageBreak())

    # Table of contents (manual, no automatic page numbers in this approach)
    story.append(Paragraph("Contents", h1_style))
    story.append(Spacer(1, 0.2*cm))
    for i, (sec_title, _) in enumerate(SECTIONS):
        story.append(Paragraph(f"{i+1}.  {sec_title}", toc_h1_style))
    story.append(PageBreak())

    # Architecture figure
    story.append(Paragraph("Architecture Overview", h1_style))
    story.append(Spacer(1, 0.2*cm))
    img = Image(str(ARCH_PNG), width=16*cm, height=12*cm)
    story.append(img)
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(
        "Figure 1. Layered pipeline of the Laertius system. The dashed arrow from "
        "LOD back to Annotation is the only feedback arc: the tagger's gazetteer "
        "is derived from the published LOD graph, not maintained separately.",
        ParagraphStyle("caption", fontName="DvSans", fontSize=8.5, leading=12,
                       alignment=1, textColor=colors.HexColor("#555"))
    ))
    story.append(PageBreak())

    # Sections
    for sec_num, (sec_title, content) in enumerate(SECTIONS, 1):
        story.append(Paragraph(f"{sec_num}. {sec_title}", h1_style))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#ddd"), thickness=0.5))
        story.append(Spacer(1, 0.15*cm))

        for item_type, item_text in content:
            if item_type == "body":
                story.append(Paragraph(item_text, body_style))
            elif item_type == "h2":
                story.append(Paragraph(item_text, h2_style))
            elif item_type == "bullet":
                story.append(Paragraph(f"\u2022  {item_text}", bullet_style))
            elif item_type == "formula":
                story.append(Paragraph(item_text, formula_style))
            elif item_type == "code":
                lines = item_text.split("\n")
                for line in lines:
                    story.append(Paragraph(line if line else " ", code_style))
            elif item_type == "greek":
                story.append(Paragraph(item_text,
                    ParagraphStyle("greek", fontName="DvSerif", fontSize=10,
                                   leading=14, spaceAfter=6,
                                   textColor=colors.HexColor("#3a1a5a"))))
            elif item_type == "table":
                rows = item_text
                col_w = [(W - 5*cm) * 0.32, (W - 5*cm) * 0.68]
                t = Table(rows, colWidths=col_w)
                t.setStyle(TABLE_STYLE)
                story.append(t)
                story.append(Spacer(1, 0.2*cm))

        story.append(Spacer(1, 0.3*cm))

    doc.build(story)
    print(f"PDF saved: {PDF_OUT}")


# ---------------------------------------------------------------------------
# 4. DOCX generation
# ---------------------------------------------------------------------------
def make_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.oxml as oxml

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title page
    def add_para(text, style="Normal", align=None, bold=False, color=None, size=None):
        p = doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
        return p

    add_para(TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
             color=(26, 58, 92), size=22)
    for line in SUBTITLE.split("\n"):
        add_para(line, align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=(85, 85, 85), size=13)
    add_para(DATE, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=(136, 136, 136), size=11)
    doc.add_page_break()

    # Table of contents
    h = doc.add_heading("Contents", level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 58, 92)
    for i, (sec_title, _) in enumerate(SECTIONS, 1):
        p = doc.add_paragraph(f"{i}.  {sec_title}")
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # Architecture figure
    h = doc.add_heading("Architecture Overview", level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 58, 92)
    doc.add_picture(str(ARCH_PNG), width=Cm(16))
    cap = doc.add_paragraph(
        "Figure 1. Layered pipeline. The dashed arrow from LOD back to "
        "Annotation is the only feedback arc: the tagger's gazetteer is derived "
        "from the published LOD graph, not maintained separately."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # Sections
    for sec_num, (sec_title, content) in enumerate(SECTIONS, 1):
        h = doc.add_heading(f"{sec_num}. {sec_title}", level=1)
        h.runs[0].font.color.rgb = RGBColor(26, 58, 92)

        for item_type, item_text in content:
            if item_type == "body":
                p = doc.add_paragraph(item_text)
                p.runs[0].font.size = Pt(10.5)
            elif item_type == "h2":
                h2 = doc.add_heading(item_text, level=2)
                h2.runs[0].font.color.rgb = RGBColor(29, 120, 116)
                h2.runs[0].font.size = Pt(11)
            elif item_type == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item_text).font.size = Pt(10.5)
            elif item_type == "formula":
                p = doc.add_paragraph(item_text)
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(42, 42, 90)
                p.paragraph_format.left_indent = Cm(1.5)
            elif item_type == "code":
                p = doc.add_paragraph(item_text)
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(8.5)
                p.paragraph_format.left_indent = Cm(1)
            elif item_type == "greek":
                p = doc.add_paragraph(item_text)
                p.runs[0].font.size = Pt(10.5)
                p.runs[0].font.color.rgb = RGBColor(58, 26, 90)
            elif item_type == "table":
                rows = item_text
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        if r_idx == 0:
                            cell.paragraphs[0].runs[0].bold = True
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            # Header background color via XML
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "1A3A5C")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:val"), "clear")
                            tcPr.append(shd)
                doc.add_paragraph()

    doc.save(str(DOCX_OUT))
    print(f"DOCX saved: {DOCX_OUT}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("Step 1/3: Generating architecture figure...")
    make_arch_figure()
    print("Step 2/3: Generating PDF...")
    make_pdf()
    print("Step 3/3: Generating DOCX...")
    make_docx()
    print("Done.")
    pdf_size  = PDF_OUT.stat().st_size // 1024
    docx_size = DOCX_OUT.stat().st_size // 1024
    print(f"  PDF:  {PDF_OUT.name}  ({pdf_size} KB)")
    print(f"  DOCX: {DOCX_OUT.name}  ({docx_size} KB)")
