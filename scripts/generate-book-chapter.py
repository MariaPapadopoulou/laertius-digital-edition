"""
Generate the book-chapter text presenting the Laertius digital edition,
focused on knowledge modelling and the production pipeline.

Produces: exports/laertius-book-chapter.docx (about three pages).
No em dashes anywhere in this file or the output document.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

WORKSPACE = Path(__file__).resolve().parent.parent
OUT = WORKSPACE / "exports" / "laertius-book-chapter.docx"

TITLE = "Modelling the Lives: Knowledge Representation and the Production Pipeline of a Digital Scholarly Edition of Diogenes Laertius"

SECTIONS = [
    (None, [
        "The Lives and Opinions of Eminent Philosophers of Diogenes Laertius is at once a biography, an anthology, and a doxography: ten books that transmit, often verbatim, the sayings, letters, wills, verses and doctrines of Greek philosophy from Thales to Epicurus. A digital scholarly edition of such a text cannot be a page-turner with a search box. What the Lives asserts, who asserts it, in which embedded document, and with what reliability: these are the questions a reader actually brings to the work, and they are questions of knowledge modelling. The edition presented here treats the text of the Lives (the Greek of the Perseus TEI transcription with the English translation of R. D. Hicks, 1925) as the fixed evidential ground, and builds above it a hand-curated knowledge graph, an ontoterminology, and a retrieval system, all regenerated from the same curated sources by a single deterministic pipeline."
    ]),
    ("The assertion as the unit of knowledge", [
        "The central modelling decision is that the atom of the edition is not the fact but the assertion. The Lives rarely states; it reports. Diogenes writes that Apollodorus places a birth in a given Olympiad, that Sosicrates disputes a succession, that \u2018some say\u2019 a philosopher died of an illness. The edition therefore curates more than fifteen hundred assertions, each carrying five obligatory coordinates: the subject it concerns, the topical content, the authority who asserts it (Diogenes himself or a named cited source such as Apollodorus or Antisthenes), the exact CTS-addressed passage that carries the evidence, and a certainty grade. Certainty is four-valued: asserted, when Diogenes speaks in his own voice; reported, when he transmits a named or anonymous source; disputed, when the text itself records disagreement; conjectured, for the rare editorial inference. Nothing in the graph floats free of the text: every assertion resolves to a section of the Lives, and the passage layer quotes the evidence verbatim in both languages, so a sceptical reader is always one click from the Greek.",
        "Around the assertions stands the second characteristic layer, the embedded microgenres. The Lives quotes whole documents inside itself, and the edition models them as first-class objects rather than as decoration: six wills, thirty-one letters, hundreds of epigrams and quoted verses, several hundred apophthegmata, anecdotes, and doxographic summaries. Each embedded document is typed (Testament, Epistle, Verse, Saying, Anecdote, Opinions), anchored to its philosopher, and carries its verbatim excerpt as a language-tagged text object. Attribution received particular care: a saying quoted inside the life of one philosopher is not necessarily his, and the curation distinguishes the owner of the chapter from the speaker in context, marking cross-attributed material explicitly."
    ]),
    ("Ontoterminology: concepts apart from words", [
        "The conceptual scaffolding follows the ontoterminology programme of Christophe Roche: a terminology whose conceptual system is an ontology, with a strict separation between the conceptual dimension and the linguistic dimension. The concept system was authored in TEDI and is mirrored one to one in the edition; the pipeline keeps the same fragment identifiers, the same isA tree, the same relation and attribute signatures as the reference export, and every departure from it is flagged as a reviewable extension. Concepts (Philosopher, PhilosophicalSchool, Testament, Assertion, Topic and their kin) exist independently of any language. Terms then denote concepts: philosopher and \u03c6\u03b9\u03bb\u03cc\u03c3\u03bf\u03c6\u03bf\u03c2 both denote the concept Philosopher, \u03b4\u03b9\u03b1\u03b8\u03ae\u03ba\u03b7 denotes Testament, and a term carries its own linguistic data, part of speech, gender, status and definition, without contaminating the concept. The catalogue also records status distinctions the sources impose: \u03c8\u03c5\u03c7\u03ae is the preferred Greek term for the soul as a doctrinal subject, while \u03c0\u03bd\u03b5\u1fe6\u03bc\u03b1, the Stoic word for the soul's substance, enters as an admitted term of the same concept.",
        "Proper names receive the same two-dimensional treatment, which matters in a corpus where homonymy is endemic; Diogenes himself closes many lives with lists of namesakes. A proper name is a linguistic surface that denotes an object, allonyms are first-class, and the Greek nominative forms are curated per bearer rather than generated, so that two Herakleides of different identity never silently exchange names. At the outer edge, the conceptual layer alone is aligned to the linked open data cloud: class-level bridges to CIDOC CRM, LAWD, FaBiO and schema.org, and several hundred curated, individually verified links to Wikidata, DBpedia, VIAF, InPhO and Pleiades. The linguistic dimension is deliberately local and carries no external mapping; concepts without an honest match remain unmapped, since a wrong link is worse than none."
    ]),
    ("Semantic annotation: binding the model back to the letter of the text", [
        "The ontoterminology and the assertion layer answer what the edition knows; semantic annotation answers where. Following the W3C Web Annotation model, every occurrence of a modelled entity in the Lives, every philosopher, school, place and work, is anchored to the text by an annotation whose target is the exact character span within its CTS-addressed section and whose body is the entity of the knowledge graph. The annotation layer multiplies the base graph severalfold, and it is what turns three separate artefacts into one edition. The connection runs in both directions. Read from the text outward, a Greek or English surface form in a passage resolves through its annotation to a graph object, from the object to the concept it instantiates, and from the concept to the terms that denote it in each language: a reader who meets \u0394\u03b9\u03b1\u03b8\u1fc6\u03ba\u03b1\u03b9 in a chapter is one hop from the concept Testament, its definition, and its whole extension of six wills. Read from the model inward, every claim of the graph and every entry of the terminological dictionaries can be traversed down to the letter of the text that licenses it. The annotations thus do for the corpus what the definitions do for the concept system: they keep the semantics honest. An ontology without anchored occurrences drifts toward wishful abstraction; a text without a conceptual layer stays a string. Bound together by annotation, the concept system is answerable to the corpus and the corpus becomes navigable by concept, and the same offsets serve the machines as well, since the annotation layer ships as standard RDF that any Web Annotation client can consume."
    ]),
    ("The expert in the loop", [
        "None of this modelling is automatic, and the edition is explicit about it. Entity recognition, attribution, homonym resolution and external linking are exactly the places where automated pipelines fail silently on an ancient corpus, so here every layer that carries a scholarly claim is authored or ratified by the domain expert. The workflow is a loop rather than a handover: the machine proposes, derives and cross-checks; the expert decides. Candidate identifications and machine-generated links are treated as untrusted until verified against the sources, and material that fails verification is excluded rather than imported with a warning. The ontoterminology makes the loop institutional: the concept system belongs to the curator's TEDI environment, the edition mirrors it, and every extension the corpus forces (a new concept, a widened relation, an added term such as \u03c8\u03c5\u03c7\u03ae) is flagged as such and routed back to the curator for review in the original tool, so the authoritative model never silently forks.",
        "The same principle shapes quality control. Validators encode the expert's judgements as executable expectations, pinning curated counts and name forms so that any later change that contradicts a past decision surfaces as a failure instead of a quiet drift. Each check carries positive controls, since a test that finds nothing must prove that it could have. And where judgement cannot be mechanised, the pipeline demands human evidence: interface checks ship with screenshots and pass criteria reviewed by a person. The edition is thus not an AI reading of Diogenes Laertius but a scholar's reading, instrumented; the machinery multiplies the reach of expertise without ever substituting for it."
    ]),
    ("Hybrid, white-box AI: closing the knowledge gap for every audience", [
        "The edition's use of artificial intelligence is best described by three qualifiers. It is hybrid, because symbolic and subsymbolic techniques are yoked together rather than substituted for one another: the knowledge graph, the ontoterminology and the SPARQL layer carry the exact, curated knowledge, while multilingual embeddings and lexical ranking carry the approximate, associative reach that lets a reader ask a question in ordinary language. It is white-box, because at no point does an opaque model stand between the reader and the evidence: a question is answered with ranked passages of the Lives rather than generated prose, a competency question displays the very query that produced its table, and every entity, term and assertion is inspectable down to its RDF. And it is expert-in-the-loop, in the sense developed above: the neural components propose and retrieve, but nothing they produce enters the edition's knowledge without the curator's signature, so the trust a reader places in the graph is trust in a scholar, not in a language model.",
        "This architecture is finally an argument about audiences, because the gap it closes is the gap in knowledge between them. The specialist arrives knowing what to ask and how: SPARQL, the RDF dumps and the linked-data alignments give full, unmediated access. The student arrives knowing the questions but not the instruments: the competency catalogue, the terminological dictionaries and the annotated text supply the instruments with their workings exposed, so that using the edition teaches the method. The general reader arrives with curiosity in natural language: hybrid retrieval meets the question where it stands and answers with the text itself, never with a paraphrase that must be taken on faith. Even the machine reader, the crawler or agent that will increasingly mediate scholarship, receives the same knowledge in citable, standards-based form. One curated model, exposed at every register from plain question to formal query, means that no audience is offered a lesser truth: the layers differ in interface, never in substance."
    ]),
    ("The pipeline: curation compiled to an edition", [
        "Everything the reader sees is compiled from curated source modules by one deterministic pipeline. The curation itself is code: typed data structures, one file per book and layer, reviewed like any other source. From these the build derives, in order, the knowledge graph (tens of thousands of triples in Turtle, RDF/XML and JSON-LD, with a full Web Annotation layer that ties every entity occurrence to its character offsets in the text), the ontoterminology exports (TEDI-compatible RDF and the self-contained HTML dictionaries of terms and proper names), the embedding index for semantic search, and the assertion store of the Legomena companion, whose own graph is reconstructed from the published dataset by SPARQL rather than copied, as a proof that the published data suffice.",
        "Retrieval is hybrid. A question put to the edition is answered by fusing lexical and dense ranking (BM25 over the corpus alongside multilingual sentence embeddings, combined by reciprocal rank fusion), and the answer is presented as ranked passages of the Lives, never as unsourced prose. The same discipline governs the scholarly interface: a catalogue of competency questions is answered live by SPARQL against the graph, with each query shown, editable and reproducible in the built-in playground, so the modelling claims of the edition are falsifiable by its own readers.",
        "What distinguishes the pipeline most, however, is the weight of verification, and it deserves description in its own right. Nearly fifty independent validators run against every change to the edition, each a small named program with a single responsibility, and a change ships only when all of them pass. They fall into recognisable families. Dataset validators guard the curated layers themselves: every claim, saying, testament, epistle, verse and doxographic summary is checked for well-formed identifiers, resolvable cross-references, and verbatim agreement between the curated excerpt and the passage of the Lives it cites, in Greek as well as in English. Consistency validators guard relations between layers: that every Greek name form still belongs to its certified bearer and never migrates between homonyms, that a person's name cannot quietly drift into the list of doctrines, that the topics of assertions stay within the controlled vocabulary. Derivation validators guard the exports: the knowledge graph is re-derived and its serialisations compared triple for triple, dictionaries and documents are re-emitted and compared by content hash rather than by timestamp, and the SPARQL of every published competency question is re-executed to prove it still returns the table the page displays. Interface validators, finally, drive a real browser against the running edition and assert not that the data exist but that the reader actually sees them: that filters filter, that links land on the cited section, that an error is shown as an error.",
        "Three habits keep this battery honest. Every validator carries positive controls, because a check that finds no violations must first demonstrate that it could have found one; a query aimed at the wrong namespace passes vacuously, and vacuous green is treated as a failure of the test, not a success of the data. Counts are pinned deliberately: when the curator admits a new concept or term, the expected totals are raised by hand in the same change, so growth is always an explicit act and never an unnoticed side effect. And validators are written after mistakes as much as before them: when review catches a misattributed saying or a stale export, the lesson is converted into a permanent check, so the same error cannot recur silently. In this sense the validation suite is the edition's institutional memory, the accumulated record of everything its editor has ever decided or repaired, kept enforceable by machine. The final artefact is a self-contained static bundle, smoke-tested end to end before it ships to an ordinary web host: the edition depends on no live infrastructure to remain citable.",
        "The result is a digital scholarly edition in the strict sense of all three words: digital, because its knowledge is modelled, queryable and machine-readable down to the single assertion; scholarly, because every layer is curated, sourced, and answerable to the expert who signs it; an edition, because every statement it makes remains an interpretable, attributable, and revisable reading of the text of Diogenes Laertius."
    ]),
]


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)

    for heading, paragraphs in SECTIONS:
        if heading:
            h = doc.add_paragraph()
            hr = h.add_run(heading)
            hr.bold = True
            hr.font.size = Pt(12)
        for text in paragraphs:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.25
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    words = sum(len(t.split()) for _, ps in SECTIONS for t in ps)
    print(f"wrote {OUT} ({words} words)")


if __name__ == "__main__":
    main()
